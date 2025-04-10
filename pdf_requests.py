from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from time import sleep
import fitz
import os

images_path = 'output-images/'

def check_path_images():
    if not os.path.exists(images_path):
        print('output-images directory not found, creating one putting the images in it.')
        convert_pdf_pages_to_images("teste-pdfs", "output-images")
    else:
        print("Directory already created.")
        pass

def convert_pdf_pages_to_images(input_folder, output_folder="output-images", zoom=4.0, image_format="png"):
    """
    Converte páginas de PDFs em imagens com alta qualidade.
    
    Args:
        input_folder (str): Pasta contendo os PDFs de entrada.
        output_folder (str): Pasta onde as imagens serão salvas (fora de "teste-pdfs").
        zoom (float): Fator de zoom para aumentar a qualidade (padrão: 4.0).
        image_format (str): Formato da imagem de saída (padrão: "png").
    """
    # Cria a pasta de saída fora de "teste-pdfs"
    os.makedirs(output_folder, exist_ok=True)

    # Lista todos os arquivos PDF na pasta de entrada
    pdf_files = [os.path.join(input_folder, f) for f in os.listdir(input_folder) if f.endswith(".pdf")]

    for pdf_file in pdf_files:
        # Nome do PDF sem extensão
        pdf_name = os.path.splitext(os.path.basename(pdf_file))[0]
        # Cria uma pasta para o PDF atual dentro de "output-images"
        pdf_output_folder = os.path.join(output_folder, pdf_name)
        os.makedirs(pdf_output_folder, exist_ok=True)

        # Abre o PDF
        pdf_document = fitz.open(pdf_file)
        total_pages = len(pdf_document)  # Número total de páginas no PDF
        padding = len(str(total_pages - 1))  # Calcula o número de dígitos necessários para zero-padding

        for page_number in range(total_pages):
            # Renderiza a página como uma imagem com alta qualidade
            page = pdf_document[page_number]
            matrix = fitz.Matrix(zoom, zoom)  # Aumenta a resolução da imagem
            pix = page.get_pixmap(matrix=matrix)

            # Salva a imagem com zero-padding no nome
            output_image_path = os.path.join(pdf_output_folder, f"{str(page_number).zfill(padding)}.{image_format}")
            pix.save(output_image_path)

        pdf_document.close()
        print(f"PDF '{pdf_file}' convertido e salvo em '{pdf_output_folder}'.")

def activation():
    check_path_images()
    chrome_options = Options()
    chrome_options.add_experimental_option("detach", True)
    driver = webdriver.Chrome(options=chrome_options)
    return driver

def Pen_to_Print(browser):
    # Realiza o login no site
    print("Iniciando login no site...")
    browser.get("https://www.pen-to-print.com/App/notes/")  # Abre a página inicial
    sleep(3)  # Aguarda o carregamento da página

    # Passo 1: Localiza e clica no botão "Log in"
    wait = WebDriverWait(browser, 10)  # Aguarda até 10 segundos
    login_button = wait.until(EC.element_to_be_clickable((By.XPATH, "//div[text()='Log in']")))
    login_button.click()
    print("Botão de login clicado e popup aberto.")

    # Passo 2: Aguarda o popup de login e preenche os campos
    email_input = wait.until(EC.presence_of_element_located((By.NAME, "email")))
    email_input.send_keys("pdfferramenta@outlook.com")

    password_input = wait.until(EC.presence_of_element_located((By.NAME, "password")))
    password_input.send_keys("Gmo2Ha5z2!rio@by@Vb22hE68yj^eQKyWR%P9BB8C!58vNmjV899oTDA8CXZ82^&")

    # Clica no botão "Login" dentro do popup
    popup_login_button = wait.until(EC.element_to_be_clickable((By.CSS_SELECTOR, ".login-button button[type='submit']")))
    popup_login_button.click()
    print("Login realizado com sucesso!")
    sleep(5)  # Aguarda o login ser processado

    # Caminho para a pasta "output-images"
    folder_path = os.path.abspath("output-images")

    # Verifica se a pasta "output-images" existe, caso contrário, cria-a
    if not os.path.exists(folder_path):
        print(f"A pasta '{folder_path}' não existe. Criando...")
        os.makedirs(folder_path)

    # Lista todas as subpastas dentro de "output-images"
    subfolders = [os.path.join(folder_path, d) for d in os.listdir(folder_path) if os.path.isdir(os.path.join(folder_path, d))]

    for subfolder in subfolders:
        print(f"Processando subpasta: {subfolder}")
        # Lista todos os arquivos PNG na subpasta atual
        png_files = [os.path.join(subfolder, f) for f in sorted(os.listdir(subfolder)) if f.endswith(".png")]

        total_files = len(png_files)

        # Verifica se a subpasta tem mais ou menos de 50 imagens
        if total_files <= 50:
            print(f"Subpasta '{subfolder}' tem {total_files} imagens. Processando todas de uma vez.")
            # Envia todas as imagens de uma vez
            for png_file in png_files:
                print(f"Enviando arquivo: {png_file}")
                try:
                    upload_input = browser.find_element(By.CSS_SELECTOR, "input[type='file']")
                    browser.execute_script("arguments[0].value = '';", upload_input)
                    upload_input.send_keys(png_file)
                    sleep(1)
                except Exception as e:
                    print(f"Erro ao enviar o arquivo {png_file}: {e}")
                    continue

            # Clica no botão "converter" após enviar todas as imagens
            try:
                button = WebDriverWait(browser, 30).until(
                    EC.element_to_be_clickable((By.CLASS_NAME, "convert-button"))
                )
                browser.execute_script("arguments[0].scrollIntoView(true);", button)
                sleep(1)
                button.click()
                print("Botão de conversão clicado. Aguardando processamento...")
                WebDriverWait(browser, 30).until(
                    EC.presence_of_element_located((By.CLASS_NAME, "scanline-cell-content"))
                )
            except Exception as e:
                print(f"Erro ao clicar no botão de conversão: {e}")
                continue

            # Navega entre as páginas e salva o conteúdo em um arquivo Word
            try:
                # Cria o arquivo Word
                doc = Document()
                page_counter = browser.find_element(By.CLASS_NAME, "page-counter").text
                current_page, total_pages = map(int, page_counter.split(" ")[1].split("/"))

                while current_page <= total_pages:
                    # Extrai o texto da página atual
                    textarea = browser.find_element(By.CLASS_NAME, "scanline-cell-content")
                    text_content = textarea.get_attribute("value")
                    if text_content.strip():
                        doc.add_paragraph(text_content)  # Adiciona o texto da página
                        doc.add_paragraph("")  # Adiciona um parágrafo vazio para separar as páginas

                    # Avança para a próxima página, se houver
                    if current_page < total_pages:
                        next_button = browser.find_element(By.XPATH, "//div[@class='scanline-arrow']/img[@alt='next page']")
                        browser.execute_script("arguments[0].click();", next_button)
                        sleep(2)  # Aguarda o carregamento da próxima página
                        page_counter = browser.find_element(By.CLASS_NAME, "page-counter").text
                        current_page, total_pages = map(int, page_counter.split(" ")[1].split("/"))
                    else:
                        break

                # Salva o arquivo Word
                output_file = os.path.join("results", f"{os.path.basename(subfolder)}.docx")
                os.makedirs("results", exist_ok=True)
                doc.save(output_file)
                print(f"Texto extraído e salvo em '{output_file}'.")
            except Exception as e:
                print(f"Erro ao navegar ou salvar o texto: {e}")
                continue

        else:
            print(f"Subpasta '{subfolder}' tem {total_files} imagens. Aplicando regra de lotes de 50.")
            # Processa os arquivos em lotes de 50
            batch_size = 50
            current_index = 0

            while current_index < total_files:
                batch_files = png_files[current_index:current_index + batch_size]

                # Envia os arquivos do lote
                for png_file in batch_files:
                    print(f"Enviando arquivo: {png_file}")
                    try:
                        upload_input = browser.find_element(By.CSS_SELECTOR, "input[type='file']")
                        browser.execute_script("arguments[0].value = '';", upload_input)
                        upload_input.send_keys(png_file)
                        sleep(1)
                    except Exception as e:
                        print(f"Erro ao enviar o arquivo {png_file}: {e}")
                        continue

                # Clica no botão "converter" após o envio do lote
                try:
                    button = WebDriverWait(browser, 30).until(
                        EC.element_to_be_clickable((By.CLASS_NAME, "convert-button"))
                    )
                    browser.execute_script("arguments[0].scrollIntoView(true);", button)
                    sleep(1)
                    button.click()
                    print("Botão de conversão clicado. Aguardando processamento...")
                    WebDriverWait(browser, 30).until(
                        EC.presence_of_element_located((By.CLASS_NAME, "scanline-cell-content"))
                    )
                except Exception as e:
                    print(f"Erro ao clicar no botão de conversão: {e}")
                    continue

                # Atualiza o índice para o próximo lote
                current_index += batch_size

        print(f"Subpasta '{subfolder}' processada com sucesso.")

    print("Processamento concluído para todas as subpastas.")
    
# Executa a função
Pen_to_Print(activation())
