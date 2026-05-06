import shutil
import pypdfium2 as pdfium
import os
import threading
import uuid
import subprocess
import tempfile
import tools
import credentials
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from PIL import Image
from docx import Document
from time import sleep


images_path = 'output-images/'
cache_selemium = r'%USERPROFILE%\.cache\selenium'

# Constants
PEN_TO_PRINT_URL = "https://www.pen-to-print.com/App/notes/"
JS_CLICK = "arguments[0].click();"

# Callback defined by `main.py` to show an attention popup on the GUI thread
# `main.py` will assign a function to this variable when initializing the GUI
show_attention_popup = None

# Holds the active Selenium driver so it can be closed from anywhere (e.g. finish_program)
_active_driver = None


def delete_images():
    # Deletes the "output-images" folder if it exists
    if os.path.exists(images_path):
        shutil.rmtree(images_path)
        print("Deleted the 'output-images' folder, wait a moment...")
    else:
        print("The 'output-images' folder does not exist.")

def check_path_images():
    if not os.path.exists(images_path):
        print('output-images directory not found, creating one and putting the images in it.')
        convert_pdf_pages_to_images("output-images")
    else:
        print("Directory already created.")

def delete_cache():
    # Performs the cache deletion process
    safe_home = os.path.normpath(os.path.abspath(os.path.expanduser('~')))
    safe_cache_path = os.path.join(safe_home, '.cache', 'selenium')

    if os.path.exists(safe_cache_path):
        try:
            shutil.rmtree(safe_cache_path)
            print(f"Deleted the 'selenium' cache folder at {safe_cache_path}, wait a moment...")
        except Exception as e:
            print(f"Failed to delete selenium cache folder '{safe_cache_path}': {e}")
    else:
        print(f"Selenium cache folder does not exist: {safe_cache_path}")


def cleanup_scoped_dirs():
    # Attempt to kill any hanging chromedriver processes to release file handles
    print("Releasing Selenium handles...")
    try:
        subprocess.run(["taskkill", "/F", "/IM", "chromedriver.exe", "/T"], capture_output=True, check=False)
    except Exception:
        pass

    temp_dir = tempfile.gettempdir()
    if not os.path.exists(temp_dir):
        return

    print(f"Using Windows Shell API to clean up 'scoped_dir' folders in {temp_dir}...")
    
    # PowerShell script to use Shell.Application COM object (mimics Explorer behavior)
    # We use -EncodedCommand to avoid any potential escaping issues with paths
    ps_code = f"""
    $shell = New-Object -ComObject Shell.Application
    $tempPath = '{temp_dir}'
    $tempFolder = $shell.NameSpace($tempPath)
    if ($tempFolder) {{
        Get-ChildItem -Path $tempPath -Filter 'scoped_dir*' -Directory | ForEach-Object {{
            $oldName = $_.Name
            $item = $tempFolder.ParseName($oldName)
            if ($item) {{
                $randomName = 'delete_me_' + [Guid]::NewGuid().ToString('N')
                try {{
                    $item.Name = $randomName
                    $newItem = $tempFolder.ParseName($randomName)
                    if ($newItem) {{
                        Remove-Item -Path $newItem.Path -Recurse -Force -ErrorAction SilentlyContinue
                        Write-Host "Successfully processed and deleted $oldName (renamed to $randomName)."
                    }}
                }} catch {{
                    Write-Host "Failed to process $oldName via Shell API."
                }}
            }}
        }}
    }}
    """
    
    try:
        # Run the powershell command
        result = subprocess.run(["powershell", "-Command", ps_code], capture_output=True, text=True, check=False)
        if result.stdout:
            print(result.stdout.strip())
        
        # Check if there are still scoped_dir folders left (for logging)
        remaining = [d for d in os.listdir(temp_dir) if d.startswith("scoped_dir") and os.path.isdir(os.path.join(temp_dir, d))]
        if not remaining:
            print("Cleanup finished. All 'scoped_dir' folders removed.")
        else:
            print(f"Cleanup finished. {len(remaining)} folders could not be removed (still in use).")
            
    except Exception as e:
        print(f"Error during Shell cleanup execution: {e}")


def _render_page(pdf_document, page_number, output_image_path, zoom):
    """Renders a single PDF page to an image file using pypdfium2."""
    page = pdf_document[page_number]
    bitmap = page.render(scale=zoom)
    pil_image = bitmap.to_pil()
    pil_image.save(output_image_path)


def _render_page_with_retry(pdf_document, page_number, output_image_path, zoom):
    """Renders a page, retrying once on failure."""
    try:
        _render_page(pdf_document, page_number, output_image_path, zoom)
        print(f"Page {page_number} saved as {output_image_path}")
    except Exception as e:
        print(f"Error rendering page {page_number}: {e}")
        print(f"Trying to re-extract page {page_number} from the PDF...")
        try:
            _render_page(pdf_document, page_number, output_image_path, zoom)
            print(f"Page {page_number} re-extracted and saved as {output_image_path} (without splitting)")
        except Exception as re_extraction_error:
            print(f"Error re-extracting page {page_number}: {re_extraction_error}")


def _check_and_rerender_image(image_file, pdf_document, zoom):
    """Re-renders an image from PDF if the file is empty or missing."""
    if not os.path.exists(image_file) or os.path.getsize(image_file) == 0:
        print(f"Error: File '{image_file}' is empty or does not exist. Trying to re-extract...")
        page_number = int(os.path.basename(image_file).split('.')[0])
        _render_page(pdf_document, page_number, image_file, zoom)
        print(f"Page {page_number} re-extracted and saved as {image_file} (without splitting)")
        return True
    return False


def _split_wide_image(image_file, image_format):
    """Splits a wide (landscape) image into left and right halves."""
    with Image.open(image_file) as img:
        width, height = img.size
        if width > height:
            print(f"Image '{image_file}' is wide. Splitting...")
            left_image = img.crop((0, 0, width // 2, height))
            right_image = img.crop((width // 2, 0, width, height))
            left_image_path = image_file.replace(f".{image_format}", f"_left.{image_format}")
            right_image_path = image_file.replace(f".{image_format}", f"_right.{image_format}")
            left_image.save(left_image_path)
            right_image.save(right_image_path)
            print(f"Split image saved as '{left_image_path}' and '{right_image_path}'")
            os.remove(image_file)
        else:
            print(f"Image '{image_file}' does not need to be split.")


def _process_pdf_images(pdf_output_folder, pdf_document, zoom, image_format):
    """Validates and splits wide images in the output subfolder."""
    print(f"Checking images in the subfolder '{pdf_output_folder}'...")
    image_files = [
        os.path.join(pdf_output_folder, f)
        for f in os.listdir(pdf_output_folder)
        if f.endswith(f".{image_format}")
    ]
    for image_file in image_files:
        try:
            if _check_and_rerender_image(image_file, pdf_document, zoom):
                continue
            _split_wide_image(image_file, image_format)
        except Exception as e:
            print(f"Error processing image '{image_file}': {e}")


def _process_pdf(pdf_file, output_folder, zoom, image_format):
    """Renders all pages of a single PDF to images in a subfolder."""
    pdf_name = os.path.splitext(os.path.basename(pdf_file))[0]
    pdf_output_folder = os.path.join(output_folder, pdf_name)
    os.makedirs(pdf_output_folder, exist_ok=True)

    pdf_document = pdfium.PdfDocument(pdf_file)
    total_pages = len(pdf_document)

    for page_number in range(total_pages):
        output_image_path = os.path.join(
            pdf_output_folder, f"{str(page_number).zfill(3)}.{image_format}"
        )
        _render_page_with_retry(pdf_document, page_number, output_image_path, zoom)

    _process_pdf_images(pdf_output_folder, pdf_document, zoom, image_format)
    pdf_document.close()


def convert_pdf_pages_to_images(output_folder="output-images", zoom=4.0, image_format="png"):
    """
    Converts PDF pages into high-quality images and checks dimensions to split wide images.
    """
    # Merge all PDFs into a single file before processing images
    tools.merge_pdfs()
    merged_pdf_path = os.path.join("./results/", "merged.pdf")

    os.makedirs(output_folder, exist_ok=True)

    pdf_files = [merged_pdf_path]

    # Create and start threads for the merged PDF file
    threads = []
    for pdf_file in pdf_files:
        thread = threading.Thread(target=_process_pdf, args=(pdf_file, output_folder, zoom, image_format))
        threads.append(thread)
        thread.start()

    for thread in threads:
        thread.join()

    print("Processing completed.")
    print("Starting the Pen-to-Print scrapper, wait a moment...")


def activation():
    global _active_driver
    check_path_images()
    chrome_options = Options()
    chrome_options.add_experimental_option("detach", True)
    driver = webdriver.Chrome(options=chrome_options)
    _active_driver = driver
    return driver


def close_browser():
    """Quit the active Selenium browser if one is open."""
    global _active_driver
    if _active_driver is not None:
        try:
            _active_driver.quit()
            print("Browser closed.")
        except Exception as e:
            print(f"Error closing browser: {e}")
        finally:
            _active_driver = None


def _login_accept_cookies(browser):
    """Clicks the 'Accept cookies' banner if present."""
    try:
        accept_button = WebDriverWait(browser, 10).until(
            EC.element_to_be_clickable((By.ID, "rcc-confirm-button"))
        )
        browser.execute_script(JS_CLICK, accept_button)
        print("'Accept cookies' button clicked successfully.")
    except Exception as e:
        print(f"Error clicking the 'Accept cookies' button: {e}")


def _get_or_request_credentials():
    """Returns (email, password) from secure storage or prompts the user."""
    try:
        email, password = credentials.get_credentials()
        print("Credentials obtained from secure storage.")
        return email, password
    except Exception:
        pass

    try:
        if callable(show_attention_popup):
            show_attention_popup()
    except Exception:
        pass

    print("Credentials not found. Enter them in the GUI console to proceed.")
    try:
        email = input("Email para login: ")
        password = input("Senha para login: ")
    except Exception:
        print("Failed to read user input. Aborting login.")
        return None, None

    if not email or not password:
        print("Email or password empty. Aborting login.")
        return None, None

    try:
        credentials.save_credentials(email, password)
        print("Credentials saved successfully.")
    except Exception as e:
        print(f"Warning: could not save credentials: {e}")

    return email, password


def _perform_login(browser):
    """Fills and submits the login form. Returns False if credentials are unavailable."""
    wait = WebDriverWait(browser, 10)
    login_button = wait.until(EC.element_to_be_clickable((By.XPATH, "//div[text()='Log in']")))
    login_button.click()
    print("Login button clicked and popup opened.")

    email, password = _get_or_request_credentials()
    if not email:
        return False

    email_input = wait.until(EC.presence_of_element_located((By.NAME, "email")))
    email_input.clear()
    email_input.send_keys(email)

    password_input = wait.until(EC.presence_of_element_located((By.NAME, "password")))
    password_input.clear()
    password_input.send_keys(password)

    popup_login_button = wait.until(
        EC.element_to_be_clickable((By.CSS_SELECTOR, ".login-button button[type='submit']"))
    )
    popup_login_button.click()
    print("Login successful!")
    sleep(5)
    return True


def _extract_single_page_text(browser, doc):
    """Extracts text from a single-page result and appends it to the Word document."""
    try:
        textarea = WebDriverWait(browser, 10).until(
            EC.presence_of_element_located((By.CLASS_NAME, "scanline-cell-content"))
        )
        text_content = textarea.get_attribute("value")
        if text_content.strip():
            doc.add_paragraph(text_content)
            doc.add_paragraph("")
    except Exception as e:
        print(f"Error scraping single-page text: {e}")


def _extract_multi_page_text(browser, doc):
    """Navigates through all result pages and appends extracted text to the Word document."""
    page_counter = browser.find_element(By.CLASS_NAME, "page-counter").text
    current_page, total_pages = map(int, page_counter.split(" ")[1].split("/"))

    while current_page <= total_pages:
        textarea = browser.find_element(By.CLASS_NAME, "scanline-cell-content")
        text_content = textarea.get_attribute("value")
        if text_content.strip():
            doc.add_paragraph(text_content)
            doc.add_paragraph("")

        if current_page < total_pages:
            next_button = browser.find_element(
                By.XPATH, "//div[@class='scanline-arrow']/img[@alt='next page']"
            )
            browser.execute_script(JS_CLICK, next_button)
            sleep(2)
            page_counter = browser.find_element(By.CLASS_NAME, "page-counter").text
            current_page, total_pages = map(int, page_counter.split(" ")[1].split("/"))
        else:
            break


def _process_batch(browser, batch_files, doc, output_file):
    """Uploads a batch of images, converts them, extracts text, and saves the document."""
    for png_file in batch_files:
        print(f"Uploading file: {png_file}")
        try:
            upload_input = browser.find_element(By.CSS_SELECTOR, "input[type='file']")
            browser.execute_script("arguments[0].value = '';", upload_input)
            upload_input.send_keys(png_file)
            sleep(1)
        except Exception as e:
            print(f"Error uploading the file {png_file}: {e}")
            continue

    # Click the convert button
    try:
        button = WebDriverWait(browser, 30).until(
            EC.element_to_be_clickable((By.CLASS_NAME, "convert-button"))
        )
        browser.execute_script("arguments[0].scrollIntoView(true);", button)
        sleep(1)
        button.click()
        print("Convert button clicked. Waiting for processing...")
        WebDriverWait(browser, 30).until(
            EC.presence_of_element_located((By.CLASS_NAME, "scanline-cell-content"))
        )
    except Exception as e:
        print(f"Error clicking the convert button: {e}")
        return

    # Extract text
    try:
        if len(batch_files) == 1:
            _extract_single_page_text(browser, doc)
        else:
            _extract_multi_page_text(browser, doc)
        doc.save(output_file)
        print(f"Text extracted and saved in '{output_file}'.")
    except Exception as e:
        print(f"Error navigating or saving the text: {e}")
        return

    # Close the result window
    try:
        close_button = WebDriverWait(browser, 10).until(
            EC.element_to_be_clickable((By.CLASS_NAME, "closeWindowButton"))
        )
        browser.execute_script(JS_CLICK, close_button)
        print("'closeWindowButton' clicked successfully.")
        sleep(2)
    except Exception as e:
        print(f"Error clicking the 'closeWindowButton': {e}")


def _process_subfolder(browser, subfolder):
    """Processes all PNG images inside a single subfolder and writes a Word document."""
    print(f"Processing subfolder: {subfolder}")
    png_files = [
        os.path.join(subfolder, f)
        for f in sorted(os.listdir(subfolder))
        if f.endswith(".png")
    ]
    total_files = len(png_files)
    output_file = os.path.join("results", f"{os.path.basename(subfolder)}.docx")
    os.makedirs("results", exist_ok=True)
    doc = Document()

    batch_size = 50
    current_index = 0

    while current_index < total_files:
        batch_files = png_files[current_index:current_index + batch_size]
        _process_batch(browser, batch_files, doc, output_file)
        current_index += batch_size

        if current_index < total_files:
            print("Restarting the process for the remaining images...")
            browser.get(PEN_TO_PRINT_URL)
            sleep(5)

    print(f"Subfolder '{subfolder}' processed successfully. Returning to the homepage...")
    browser.get(PEN_TO_PRINT_URL)
    sleep(5)


def pen_to_print(browser):
    """Logs into Pen-to-Print and converts all images in output-images to Word documents."""
    print("Starting login on the website...")
    browser.get(PEN_TO_PRINT_URL)
    sleep(3)

    _login_accept_cookies(browser)

    if not _perform_login(browser):
        return

    folder_path = os.path.abspath("output-images")
    if not os.path.exists(folder_path):
        print(f"The folder '{folder_path}' does not exist. Creating...")
        os.makedirs(folder_path)

    subfolders = [
        os.path.join(folder_path, d)
        for d in os.listdir(folder_path)
        if os.path.isdir(os.path.join(folder_path, d))
    ]

    for subfolder in subfolders:
        _process_subfolder(browser, subfolder)

    print("Processing completed for all subfolders.")
    delete_images()
    browser.quit()
    delete_cache()
    cleanup_scoped_dirs()